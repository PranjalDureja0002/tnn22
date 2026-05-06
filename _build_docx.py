"""Builds the PwC-styled DOCX from Agentic_Portal_Solution_Design.md."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PWC_ORANGE = RGBColor(0xD0, 0x4A, 0x02)
PWC_ORANGE_HEX = 'D04A02'
PWC_BURGUNDY = RGBColor(0x5F, 0x23, 0x35)
PWC_BURGUNDY_HEX = '5F2335'
DARK_GREY = RGBColor(0x40, 0x40, 0x40)
LIGHT_GREY_HEX = 'F5F5F5'
HEADER_GREY = RGBColor(0x80, 0x80, 0x80)


def add_horizontal_line(paragraph, color_hex=PWC_ORANGE_HEX, size=12):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def shade_paragraph(paragraph, color_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    pPr.append(shd)


def setup_styles(doc):
    s = doc.styles['Normal']
    s.font.name = 'Arial'
    s.font.size = Pt(10)
    s.paragraph_format.space_after = Pt(6)
    s.paragraph_format.line_spacing = 1.15

    h = doc.styles['Heading 1']
    h.font.name = 'Arial'
    h.font.size = Pt(18)
    h.font.bold = True
    h.font.color.rgb = PWC_BURGUNDY
    h.paragraph_format.space_before = Pt(18)
    h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.keep_with_next = True

    h = doc.styles['Heading 2']
    h.font.name = 'Arial'
    h.font.size = Pt(14)
    h.font.bold = True
    h.font.color.rgb = PWC_ORANGE
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(4)
    h.paragraph_format.keep_with_next = True

    h = doc.styles['Heading 3']
    h.font.name = 'Arial'
    h.font.size = Pt(11)
    h.font.bold = True
    h.font.color.rgb = DARK_GREY
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(3)
    h.paragraph_format.keep_with_next = True


INLINE_PATTERN = re.compile(
    r'(\*\*[^*]+?\*\*|\*[^*]+?\*|`[^`]+?`|\[[^\]]+?\]\([^)]+?\))'
)


def add_inline_runs(paragraph, text, base_italic=False):
    if not text:
        return
    parts = INLINE_PATTERN.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            if base_italic:
                run.italic = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            if base_italic:
                run.italic = True
        elif part.startswith('['):
            m = re.match(r'\[([^\]]+)\]\(([^)]+)\)', part)
            if m:
                run = paragraph.add_run(m.group(1))
                run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
                run.underline = True
            else:
                run = paragraph.add_run(part)
                if base_italic:
                    run.italic = True
        else:
            run = paragraph.add_run(part)
            if base_italic:
                run.italic = True


def split_table_row(row):
    cells = row.strip().split('|')
    if cells and not cells[0].strip():
        cells = cells[1:]
    if cells and not cells[-1].strip():
        cells = cells[:-1]
    return [c.strip() for c in cells]


def is_separator_row(line):
    cells = split_table_row(line)
    return all(re.match(r'^:?-{2,}:?$', c) for c in cells) if cells else False


def add_table_from_markdown(doc, table_lines):
    if len(table_lines) < 2:
        return
    headers = split_table_row(table_lines[0])
    if not headers:
        return
    sep_idx = 1 if is_separator_row(table_lines[1]) else None
    body_start = (sep_idx + 1) if sep_idx is not None else 1
    rows = [split_table_row(l) for l in table_lines[body_start:]]
    rows = [r for r in rows if r]

    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    table.autofit = True

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        shade_cell(cell, PWC_BURGUNDY_HEX)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for r, row in enumerate(rows):
        zebra = r % 2 == 1
        for c, val in enumerate(row):
            if c >= len(headers):
                continue
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            if zebra:
                shade_cell(cell, 'FAF7F4')
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            add_inline_runs(p, val)
            for run in p.runs:
                if run.font.size is None:
                    run.font.size = Pt(9)
                if run.font.name is None or run.font.name == 'Arial':
                    run.font.name = 'Arial'

    set_table_borders(table)


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'CCCCCC')
        tblBorders.append(b)
    tblPr.append(tblBorders)


def add_code_block(doc, code):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.left_indent = Cm(0.3)
    shade_paragraph(p, LIGHT_GREY_HEX)
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)


def render_markdown_to_doc(doc, md_text):
    lines = md_text.split('\n')
    i = 0
    n = len(lines)
    skip_first_h1 = True

    while i < n:
        line = lines[i]
        stripped = line.rstrip()
        lstripped = line.lstrip()

        if stripped.startswith('# ') and not stripped.startswith('## '):
            if skip_first_h1:
                skip_first_h1 = False
                i += 1
                continue
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=2)
        elif stripped.startswith('#### '):
            doc.add_heading(stripped[5:], level=3)
        elif stripped.startswith('```'):
            code_lines = []
            i += 1
            while i < n and not lines[i].rstrip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, '\n'.join(code_lines))
        elif lstripped.startswith('|'):
            table_lines = []
            while i < n and lines[i].lstrip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            add_table_from_markdown(doc, table_lines)
            continue
        elif stripped == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(8)
            add_horizontal_line(p, 'CCCCCC', 4)
        elif lstripped.startswith('> '):
            quote_lines = []
            while i < n and lines[i].lstrip().startswith('>'):
                cur = lines[i].lstrip()
                quote_lines.append(cur[2:] if cur.startswith('> ') else cur[1:].lstrip())
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.7)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '24')
            left.set(qn('w:space'), '8')
            left.set(qn('w:color'), PWC_ORANGE_HEX)
            pBdr.append(left)
            pPr.append(pBdr)
            add_inline_runs(p, ' '.join(quote_lines), base_italic=True)
            continue
        elif lstripped.startswith('- ') or lstripped.startswith('* '):
            while i < n and (lines[i].lstrip().startswith('- ') or lines[i].lstrip().startswith('* ')):
                bullet_text = lines[i].lstrip()[2:]
                p = doc.add_paragraph(style='List Bullet')
                add_inline_runs(p, bullet_text)
                i += 1
            continue
        elif re.match(r'^\s*\d+\.\s', stripped):
            while i < n and re.match(r'^\s*\d+\.\s', lines[i]):
                num_text = re.sub(r'^\s*\d+\.\s', '', lines[i])
                p = doc.add_paragraph(style='List Number')
                add_inline_runs(p, num_text)
                i += 1
            continue
        elif stripped == '':
            pass
        else:
            para_lines = [stripped]
            while i + 1 < n:
                nxt = lines[i + 1]
                nxt_l = nxt.lstrip()
                if (
                    nxt.rstrip() == ''
                    or nxt.startswith('#')
                    or nxt_l.startswith('|')
                    or nxt.startswith('```')
                    or nxt_l.startswith('- ')
                    or nxt_l.startswith('* ')
                    or nxt_l.startswith('> ')
                    or re.match(r'^\s*\d+\.\s', nxt)
                    or nxt.rstrip() == '---'
                ):
                    break
                i += 1
                para_lines.append(lines[i].rstrip())
            p = doc.add_paragraph()
            add_inline_runs(p, ' '.join(para_lines))

        i += 1


def create_cover_page(doc):
    line_p = doc.add_paragraph()
    line_p.paragraph_format.space_after = Pt(0)
    add_horizontal_line(line_p, PWC_ORANGE_HEX, 18)

    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(6)
    run = title_p.add_run('Agentic Portal')
    run.font.name = 'Garamond'
    run.font.size = Pt(44)
    run.font.italic = True
    run.font.color.rgb = PWC_BURGUNDY

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(0)
    run = sub_p.add_run('Solution Design Document')
    run.font.name = 'Garamond'
    run.font.size = Pt(22)
    run.font.color.rgb = PWC_ORANGE

    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    metadata = [
        'Document version 0.2 (Draft for review)',
        'Issued May 2026',
        'Classification: Internal — Confidential',
        'Prepared by PwC India — Agentic Automation Practice',
    ]
    for line in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(line)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_GREY

    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    desc_p = doc.add_paragraph()
    desc_p.paragraph_format.space_after = Pt(0)
    desc_p.paragraph_format.line_spacing = 1.4
    run = desc_p.add_run(
        "This document is the canonical solution design for the Agentic Portal — PwC India's "
        "internal practice-wide knowledge and reuse platform for the Agentic Automation Practice. "
        "It covers the problem statement, personas and journeys, information architecture, the "
        "role + grade access model, solution architecture (Azure SQL + ChromaDB + Azure OpenAI on "
        "an open-source stack), the Phase 1 AI feature set, end-to-end flows, the data model, "
        "security and compliance posture, the phased roadmap, risks, success metrics, and open "
        "questions for practice leadership and IT."
    )
    run.font.name = 'Arial'
    run.font.size = Pt(10)


def configure_body_section(doc):
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    new_section.top_margin = Cm(2.5)
    new_section.bottom_margin = Cm(2.5)
    new_section.left_margin = Cm(2.5)
    new_section.right_margin = Cm(2.5)

    new_section.header.is_linked_to_previous = False
    new_section.footer.is_linked_to_previous = False

    header = new_section.header
    header_p = header.paragraphs[0]
    header_p.text = ''
    add_horizontal_line(header_p, PWC_ORANGE_HEX, 8)
    run = header_p.add_run('Agentic Portal — Solution Design Document')
    run.font.name = 'Arial'
    run.font.size = Pt(8)
    run.font.color.rgb = HEADER_GREY

    footer = new_section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    page_run = footer_p.add_run()
    page_run.font.name = 'Arial'
    page_run.font.size = Pt(8)
    page_run.font.color.rgb = HEADER_GREY

    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'PAGE   \\* MERGEFORMAT'
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    page_run._r.append(fld_begin)
    page_run._r.append(instr)
    page_run._r.append(fld_end)


def main():
    md_path = Path(__file__).parent / 'Agentic_Portal_Solution_Design.md'
    out_path = md_path.with_suffix('.docx')

    md_text = md_path.read_text(encoding='utf-8')

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.different_first_page_header_footer = True

    setup_styles(doc)
    create_cover_page(doc)
    configure_body_section(doc)
    render_markdown_to_doc(doc, md_text)

    try:
        doc.save(out_path)
    except PermissionError:
        # File is locked (probably open in Word) — write a timestamped fallback.
        from datetime import datetime
        stamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        out_path = md_path.with_name(f'{md_path.stem}.{stamp}.docx')
        doc.save(out_path)
        print('NOTE: original DOCX was locked (open in Word?); wrote a new file instead.')
    print(f'Wrote {out_path} ({out_path.stat().st_size:,} bytes)')


if __name__ == '__main__':
    main()
